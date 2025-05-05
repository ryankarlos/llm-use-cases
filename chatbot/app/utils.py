import io
from PyPDF2 import PdfReader
import docx
import pandas as pd


def extract_text_from_file(uploaded_file):
    """Extract text content from various file types"""
    text = ""
    file_type = uploaded_file.type
    file_info = {}

    try:
        # Text files
        if file_type == "text/plain":
            text = uploaded_file.getvalue().decode("utf-8")
            file_info = {
                "type": "Text file",
                "details": f"{len(text.splitlines())} lines",
            }

        # PDF files
        elif file_type == "application/pdf":
            reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
            page_count = len(reader.pages)
            file_info = {"type": "PDF document", "details": f"{page_count} pages"}
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                text += f"--- Page {page_num} ---\n{page_text}\n\n"

        # Word documents
        elif (
            file_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            paragraph_count = len(doc.paragraphs)
            file_info = {
                "type": "Word document",
                "details": f"{paragraph_count} paragraphs",
            }
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    text += para.text + "\n"

        # CSV files
        elif file_type == "text/csv":
            try:
                df = pd.read_csv(io.BytesIO(uploaded_file.getvalue()))
            except UnicodeDecodeError:
                df = pd.read_csv(
                    io.BytesIO(uploaded_file.getvalue()), encoding="ISO-8859-1"
                )
            row_count = len(df)
            col_count = len(df.columns)
            file_info = {
                "type": "CSV spreadsheet",
                "details": f"{row_count} rows × {col_count} columns",
            }
            text = f"CSV File with {row_count} rows and {col_count} columns.\n"
            text += f"Column names: {', '.join(df.columns)}\n\n"
            text += "First 5 rows:\n" + df.head().to_string() + "\n\n"
            if row_count > 5:
                text += f"... and {row_count - 5} more rows"

        # Excel files
        elif file_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ]:
            df = pd.read_excel(io.BytesIO(uploaded_file.getvalue()))
            row_count = len(df)
            col_count = len(df.columns)
            file_info = {
                "type": "Excel spreadsheet",
                "details": f"{row_count} rows × {col_count} columns",
            }
            text = f"Excel File with {row_count} rows and {col_count} columns.\n"
            text += f"Column names: {', '.join(df.columns)}\n\n"
            text += "First 5 rows:\n" + df.head().to_string() + "\n\n"
            if row_count > 5:
                text += f"... and {row_count - 5} more rows"

        else:
            text = f"Unsupported file type: {file_type}"
            file_info = {"type": "Unsupported file", "details": file_type}

    except Exception as e:
        text = f"Error processing file: {str(e)}"
        file_info = {"type": "Error", "details": str(e)}

    return text, file_info
